# --- Importaciones necesarias de Django y otras librerías externas ---
from django.shortcuts import render, redirect  # Para renderizar plantillas y redireccionar usuarios
from django.contrib.auth.decorators import login_required  # Para proteger vistas que requieren login
from django.contrib.auth import login, logout  # Para manejar inicio/cierre de sesión
from django.contrib.auth.forms import AuthenticationForm, UserCreationForm  # Formularios integrados de login y registro
from django.contrib import messages  # Para mostrar mensajes temporales al usuario

# Importaciones internas de la app
from .forms import ProblemaForm  # Formulario personalizado para crear problemas
from .models import Problema  # Modelo que representa un problema creado por el usuario

# Bibliotecas matemáticas y gráficas
import plotly.graph_objs as go  # Para construir figuras en Plotly
import plotly.offline as opy  # Para insertar el gráfico como HTML
from sympy import symbols, Eq, solve, lambdify  # Para manipulación simbólica
import numpy as np  # Para cálculos numéricos

# Librerías para exportación de archivos
import io  # Para manejar streams de memoria (PDF)
from reportlab.pdfgen import canvas  # Para crear PDF
from reportlab.lib.pagesizes import letter  # Tamaño de página estándar
from django.http import FileResponse, HttpResponse  # Para devolver archivos al navegador
from openpyxl import Workbook  # Para crear archivos Excel
from docx import Document  # Para crear archivos Word
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from io import BytesIO
from reportlab.pdfgen import canvas
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from django.http import HttpResponse
from .models import Problema
from .utils import resolver_problema
from .utils import generar_figura_plotly
import base64

# Vista principal: página de inicio
def index(request):
    return render(request, 'grafico/index.html')

from django.shortcuts import get_object_or_404

#editar 
@login_required
def editar_problema(request, problema_id):
    problema = get_object_or_404(Problema, id=problema_id, usuario=request.user)

    if request.method == 'POST':
        form = ProblemaForm(request.POST, instance=problema)
        if form.is_valid():
            form.save()
            messages.success(request, "Problema editado correctamente.")
            return redirect('resultado', problema_id=problema.id)
    else:
        form = ProblemaForm(instance=problema)

    return render(request, 'grafico/editar.html', {'form': form, 'problema': problema})

#duplicar
@login_required
def duplicar_problema(request, problema_id):
    problema = get_object_or_404(Problema, id=problema_id, usuario=request.user)
    nuevo_problema = Problema.objects.create(
        usuario=request.user,
        objetivo=problema.objetivo,
        funcion_objetivo=problema.funcion_objetivo,
        restricciones=problema.restricciones,
        limites=problema.limites
    )
    messages.success(request, "Problema duplicado correctamente.")
    return redirect('resultado', problema_id=nuevo_problema.id)

# Vista para crear un nuevo problema (requiere login)
@login_required
def crear_problema(request):
 
    #Permite al usuario crear un nuevo problema ingresando una función objetivo
    #y restricciones mediante un formulario.
  
    if request.method == 'POST':
        form = ProblemaForm(request.POST)
        if form.is_valid():
            problema = form.save(commit=False)
            problema.usuario = request.user
            problema.save()
            return redirect('resultado', problema_id=problema.id)
    else:
        form = ProblemaForm()
    return render(request, 'grafico/crear.html', {'form': form})

@login_required
def resultado(request, problema_id):
    problema = get_object_or_404(Problema, id=problema_id, usuario=request.user)

    try:
        optimo, puntos, expr, restricciones, pasos_restricciones, pasos_detallados, mensaje_extra = resolver_problema(
            problema.funcion_objetivo,
            problema.restricciones,
            problema.objetivo
        )
    except ValueError as e:
        # Mostrar mensaje de error directamente en la misma vista
        return render(request, 'grafico/resultado.html', {
            'problema': problema,
            'grafico': None,
            'optimo': None,
            'pasos_restricciones': [],
            'pasos_detallados': [],
            'mensaje_extra': f"❌ {e}",
            'grafica_base64': None,
        })

    # Si no hubo errores, se genera el gráfico como de costumbre
    fig = go.Figure()
    x, y = symbols('x y')

    all_x = [p[0] for p in puntos]
    all_y = [p[1] for p in puntos]
    if optimo:
        all_x.append(optimo[0][0])
        all_y.append(optimo[0][1])

    min_x = min(all_x) - 5 if all_x else 0
    max_x = max(all_x) + 5 if all_x else 25
    min_y = min(all_y) - 5 if all_y else 0
    max_y = max(all_y) + 5 if all_y else 25
    x_vals = np.linspace(min_x, max_x, 400)

    for i, (lhs, op, rhs) in enumerate(restricciones):
        try:
            eq = Eq(lhs, rhs)
            sol = solve(eq, y)
            if sol:
                f_restriccion = lambdify(x, sol[0], modules=["numpy"])
                y_vals = f_restriccion(x_vals)
                fig.add_trace(go.Scatter(
                    x=x_vals,
                    y=y_vals,
                    mode='lines',
                    line=dict(dash='dot', color='gray'),
                    name=f'Restricción {i+1}'
                ))
        except Exception as e:
            print("No se pudo graficar restricción:", e)

    if puntos:
        fig.add_trace(go.Scatter(
            x=[p[0] for p in puntos],
            y=[p[1] for p in puntos],
            mode='markers',
            marker=dict(color='blue', size=6),
            name='Puntos factibles'
        ))

        if optimo:
            puntos_ordenados = sorted(puntos, key=lambda p: (
                np.arctan2(p[1] - optimo[0][1], p[0] - optimo[0][0])
            ))
            fig.add_trace(go.Scatter(
                x=[p[0] for p in puntos_ordenados],
                y=[p[1] for p in puntos_ordenados],
                fill='toself',
                fillcolor='rgba(0, 100, 80, 0.2)',
                line=dict(color='rgba(0,0,0,0)'),
                hoverinfo='skip',
                showlegend=True,
                name='Región factible'
            ))

    if optimo:
        fig.add_trace(go.Scatter(
            x=[optimo[0][0]],
            y=[optimo[0][1]],
            mode='markers+text',
            marker=dict(color='red', size=10),
            text=[f'Óptimo\nZ = {optimo[1]}'],
            textposition='top center',
            name='Punto óptimo'
        ))

        try:
            eq_nivel = Eq(expr, optimo[1])
            sol = solve(eq_nivel, y)
            if sol:
                f_linea = lambdify(x, sol[0], modules=["numpy"])
                y_vals = f_linea(x_vals)
                fig.add_trace(go.Scatter(
                    x=x_vals,
                    y=y_vals,
                    mode='lines',
                    line=dict(color='blue', dash='dash'),
                    name=f'Línea Z = {optimo[1]}'
                ))
        except Exception as e:
            print("No se pudo graficar la línea de nivel:", e)

    fig.update_layout(
        title='Gráfica del Método Gráfico',
        xaxis_title='X',
        yaxis_title='Y',
        showlegend=True
    )
    fig.update_xaxes(range=[min_x, max_x])
    fig.update_yaxes(range=[min_y, max_y])

    imagen_bytes = fig.to_image(format="png", width=800, height=600, engine="kaleido")
    imagen_base64 = base64.b64encode(imagen_bytes).decode("utf-8")

    grafico = opy.plot(fig, auto_open=False, output_type='div')

    return render(request, 'grafico/resultado.html', {
        'problema': problema,
        'grafico': grafico,
        'optimo': optimo,
        'pasos_restricciones': pasos_restricciones,
        'pasos_detallados': pasos_detallados,
        'mensaje_extra': mensaje_extra,
        'grafica_base64': imagen_base64,
    })


@login_required
def exportar_pdf(request, problema_id):
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from PIL import Image
    from io import BytesIO
    import tempfile

    problema = get_object_or_404(Problema, id=problema_id, usuario=request.user)
    optimo, puntos, expr, restricciones, pasos_despeje, pasos_detallados, mensaje_extra = resolver_problema(
        problema.funcion_objetivo, problema.restricciones, problema.objetivo
    )

    # Generar imagen base64 de la gráfica
    fig = generar_figura_plotly(problema, optimo, puntos, restricciones, expr)
    imagen_bytes = fig.to_image(format="png", width=800, height=600, engine="kaleido")

    # Crear PDF
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    margin = 50
    y = height - margin

    def write_line(text, font="Helvetica", size=11, leading=15):
        nonlocal y
        if y < margin:
            p.showPage()
            y = height - margin
        p.setFont(font, size)
        p.drawString(margin, y, text)
        y -= leading

    # Contenido principal
    p.setFont("Helvetica-Bold", 16)
    p.drawString(margin, y, "Resultado del Método Gráfico")
    y -= 30

    write_line(f"Función objetivo: Z = {problema.funcion_objetivo}")
    write_line("Restricciones:")
    for r in problema.restricciones.splitlines():
        write_line(f"• {r}")

    if optimo:
        write_line(f"Óptimo: Z = {optimo[1]:.2f} en ({optimo[0][0]:.2f}, {optimo[0][1]:.2f})", font="Helvetica-Bold")
    else:
        write_line("No se encontró solución óptima factible.", font="Helvetica-Bold")

    # Paso a paso
    write_line("")
    write_line("Paso a paso detallado", font="Helvetica-Bold", size=13)
    for paso in pasos_detallados:
        for linea in paso.split('\n'):
            write_line(linea.strip())

    # Agregar imagen
    p.showPage()
    img_stream = BytesIO(imagen_bytes)
    with Image.open(img_stream) as img:
        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        img.save(temp.name)
        p.drawInlineImage(temp.name, margin, 150, width=500, height=300)

    p.showPage()
    p.save()
    buffer.seek(0)

    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="resultado_{problema.id}.pdf"'
    return response

#exporta en excel 
@login_required
def exportar_excel(request, problema_id):
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.drawing.image import Image as ExcelImage
    from PIL import Image
    from io import BytesIO
    import os

    problema = get_object_or_404(Problema, id=problema_id, usuario=request.user)
    optimo, puntos, expr, restricciones, pasos_despeje, pasos_detallados, mensaje_extra = resolver_problema(
        problema.funcion_objetivo, problema.restricciones, problema.objetivo
    )

    # Generar la imagen de la gráfica
    fig = generar_figura_plotly(problema, optimo, puntos, restricciones, expr)
    imagen_bytes = fig.to_image(format="png", width=800, height=600, engine="kaleido")

    # Guardar temporalmente
    temp_path = "temp_excel_plot.png"
    with open(temp_path, "wb") as f:
        f.write(imagen_bytes)

    wb = Workbook()
    ws = wb.active
    ws.title = "Resultado"

    #Estilos reutilizables 
    titulo_fill = PatternFill(start_color="FCE4EC", end_color="FCE4EC", fill_type="solid")
    titulo_font = Font(bold=True, color="000000")
    center = Alignment(horizontal="left", vertical="top")

    def write_titulo(texto):
        c = ws.cell(row=ws.max_row + 1, column=1, value=texto)
        c.font = titulo_font
        c.fill = titulo_fill
        c.alignment = center
        ws.append([])

    #Contenido 
    write_titulo("Función Objetivo")
    ws.append([problema.funcion_objetivo])
    ws.append([])

    write_titulo("Restricciones")
    for restriccion in problema.restricciones.strip().splitlines():
        ws.append([f"• {restriccion}"])
    ws.append([])

    if optimo:
        write_titulo("Resultado Óptimo")
        ws.append(["X", "Y", "Z"])
        ws.append([
            float(optimo[0][0]),
            float(optimo[0][1]),
            float(optimo[1])
        ])
    else:
        ws.append(["No se encontró solución óptima."])
    ws.append([])

    write_titulo("Paso a paso del Método Gráfico")
    for paso in pasos_detallados:
        ws.append([paso])
    ws.append([])

    #Insertar imagen
    img = ExcelImage(temp_path)
    img.anchor = f"E2"
    ws.add_image(img)

    #Preparar respuesta
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="resultado_{problema.id}.xlsx"'
    wb.save(response)

    # Limpiar imagen temporal
    if os.path.exists(temp_path):
        os.remove(temp_path)

    return response

#exporta en word
@login_required
def exportar_word(request, problema_id):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from PIL import Image
    from io import BytesIO

    problema = get_object_or_404(Problema, id=problema_id, usuario=request.user)
    optimo, puntos, expr, restricciones, pasos_despeje, pasos_detallados, mensaje_extra = resolver_problema(
        problema.funcion_objetivo, problema.restricciones, problema.objetivo
    )

    # Generar la gráfica como imagen
    fig = generar_figura_plotly(problema, optimo, puntos, restricciones, expr)
    imagen_bytes = fig.to_image(format="png", width=800, height=600, engine="kaleido")

    doc = Document()

    #Estilo genera
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    def agregar_titulo(texto, nivel=1):
        p = doc.add_paragraph()
        run = p.add_run(texto)
        run.bold = True
        run.font.size = Pt(14 if nivel == 1 else 12)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.space_after = Pt(6)

    # Contenido
    agregar_titulo("Resultado del Problema", nivel=1)

    agregar_titulo("Función Objetivo", nivel=2)
    doc.add_paragraph(f"Z = {problema.funcion_objetivo}")

    agregar_titulo("Restricciones", nivel=2)
    for restriccion in problema.restricciones.splitlines():
        doc.add_paragraph(f"• {restriccion}")

    if optimo:
        agregar_titulo("Resultado Óptimo", nivel=2)
        doc.add_paragraph(f"Z = {optimo[1]:.2f} en (x, y) = ({optimo[0][0]:.2f}, {optimo[0][1]:.2f})")
    else:
        doc.add_paragraph("No se encontró solución óptima factible.")

    agregar_titulo("Paso a paso detallado", nivel=2)
    for paso in pasos_detallados:
        doc.add_paragraph(paso)

    # Insertar imagen directamente desde memoria
    with BytesIO() as temp_img:
        Image.open(BytesIO(imagen_bytes)).save(temp_img, format="PNG")
        temp_img.seek(0)
        doc.add_picture(temp_img, width=Inches(5.5))

    # Exportar archivo
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="resultado_{problema.id}.docx"'
    return response


# Vista del historial del usuario
@login_required
def historial(request):

    #Muestra todos los problemas creados por el usuario en orden descendente.

    problemas = Problema.objects.filter(usuario=request.user).order_by('-creado')
    return render(request, 'grafico/historial.html', {'problemas': problemas})


# Borrar todos los problemas del historial
@login_required
def borrar_historial(request):
    #Elimina todos los problemas del historial del usuario actual.

    if request.method == "POST":
        Problema.objects.filter(usuario=request.user).delete()
    return redirect('historial')


# Cargar un ejemplo automático
@login_required
def ejemplo(request):
    #Crea automáticamente un problema de ejemplo para mostrar funcionalidad.

    ejemplo = Problema.objects.create(
        usuario=request.user,
        objetivo='max',
        funcion_objetivo='3*x + 2*y',
        restricciones='2*x + y <= 18\n2*x + 3*y <= 42\n3*x + y <= 24',
        limites='x>=0, y>=0'
    )
    return redirect('resultado', problema_id=ejemplo.id)


# Vista de ayuda general
def ayuda(request):
   #Muestra una página estática de ayuda al usuario.

    return render(request, 'grafico/ayuda.html')


from django.shortcuts import render, redirect
from django.contrib.auth import login
from .forms import RegistroFormPersonalizado

# Vista de registro de usuario
def registro(request):
    #Permite registrar un nuevo usuario utilizando un formulario personalizado y autenticarlo automáticamente.
    
    if request.method == 'POST':
        form = RegistroFormPersonalizado(request.POST)
        if form.is_valid():
            user = form.save(commit=False)

            # Guardar campos personalizados antes de guardar el usuario
            user.first_name = form.cleaned_data['first_name']
            user.last_name = form.cleaned_data['last_name']
            user.email = form.cleaned_data['email']
            user.save()

            login(request, user)
            return redirect('index')
    else:
        form = RegistroFormPersonalizado()

    return render(request, 'grafico/registro.html', {'form': form})

# Vista de inicio de sesión
def login_view(request):
#Permite a los usuarios autenticarse en el sistema.
  
    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            login(request, form.get_user())
            next_url = request.GET.get('next') or 'index'
            return redirect(next_url)
    else:
        form = AuthenticationForm()
    return render(request, 'grafico/login.html', {'form': form})

# Cerrar sesión
def logout_view(request):
    #Cierra la sesión del usuario actual.

    logout(request)
    return redirect('index')
